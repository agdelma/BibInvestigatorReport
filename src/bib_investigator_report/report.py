r"""Generate a LaTeX report of BibTeX citations per investigator per project-year.

Inputs:
  1) BibTeX (.bib) file
  2) Investigators CSV (at minimum: Firstname, Lastname)
  3) --start-date=YYYY-MM-DD
  4) --end-date=YYYY-MM-DD

Output:
  - A .tex file organized by investigator and project-year.
  - A summary_pubs.csv file (investigators x years) with BibTeX keys per cell.

Project-year definition:
  Year 01 is the 12-month period starting at --start-date.
  Year 02 is the next 12-month period, etc, until --end-date.

Inclusion logic:
  A BibTeX entry is included for an investigator in a project-year if:
    - the investigator's name matches an author in the BibTeX 'author' field, AND
    - the inferred entry date falls inside that project-year.

Date inference:
  BibTeX often has only a 'year'. This script tries, in order:
    1) 'date' field (ISO formats: YYYY, YYYY-MM, YYYY-MM-DD)
    2) 'year' + 'month' fields
    3) 'date-added' field (common in Zotero exports)
    4) fallback: mid-year (July 1) for the given 'year'

You can force a particular field via --date-field.

LaTeX generation:
  - Uses \usepackage{multibib}.
  - Creates one bibliography per project-year via \newcites{one}{Year 01}, ...
  - Color-codes each citation by how many investigators from your CSV are on that paper:
      1 investigator: black
      2 investigators: blue
      3+ investigators: orange

PDF compilation:
  If --compile is set, the script runs pdflatex + bibtex (for each multibib aux) + pdflatex twice.

Notes on BibTeX parsing:
  This script includes a lightweight, dependency-free BibTeX parser. If you have
  pybtex or bibtexparser installed, it can be extended, but this version does not
  require external libraries.
"""

from __future__ import annotations

import argparse
import csv
import html
import datetime as dt
import re
import shutil
import subprocess
import sys
import unicodedata
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional, Sequence, Tuple
try:
    from docx import Document
except ImportError:
    Document = None


# ----------------------------
# Utilities: dates
# ----------------------------


def parse_iso_date(s: str) -> Optional[dt.date]:
    """Parse YYYY, YYYY-MM, or YYYY-MM-DD into a date.

    - YYYY -> July 1 of that year
    - YYYY-MM -> 15th of that month
    - YYYY-MM-DD -> exact
    """
    s = s.strip()
    if not s:
        return None
    s = s.replace("/", "-")

    m = re.fullmatch(r"(\d{4})", s)
    if m:
        y = int(m.group(1))
        return dt.date(y, 7, 1)

    m = re.fullmatch(r"(\d{4})-(\d{1,2})", s)
    if m:
        y, mo = int(m.group(1)), int(m.group(2))
        mo = max(1, min(12, mo))
        return dt.date(y, mo, 15)

    m = re.fullmatch(r"(\d{4})-(\d{1,2})-(\d{1,2})", s)
    if m:
        y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
        d = max(1, min(31, d))
        try:
            return dt.date(y, mo, d)
        except ValueError:
            # Clamp (e.g. Feb 30 -> Feb 28)
            while d > 1:
                d -= 1
                try:
                    return dt.date(y, mo, d)
                except ValueError:
                    continue
    return None


def add_years(d: dt.date, years: int) -> dt.date:
    """Add years to a date, clamping Feb 29 to Feb 28 as needed."""
    try:
        return d.replace(year=d.year + years)
    except ValueError:
        return d.replace(month=2, day=28, year=d.year + years)


@dataclass(frozen=True)
class Period:
    index: int
    start: dt.date  # inclusive
    end: dt.date    # exclusive

    @property
    def label(self) -> str:
        return f"Year {self.index:02d}"


def build_periods(start: dt.date, end_inclusive: dt.date) -> List[Period]:
    periods: List[Period] = []
    idx = 1
    cur = start
    end_exclusive = end_inclusive + dt.timedelta(days=1)
    while cur < end_exclusive:
        nxt = add_years(cur, 1)
        periods.append(Period(index=idx, start=cur, end=min(nxt, end_exclusive)))
        cur = nxt
        idx += 1
    return periods


# ----------------------------
# Utilities: BibTeX parsing
# ----------------------------


def _split_top_level(s: str, sep: str = ",") -> List[str]:
    """Split on sep at brace depth 0 and outside quotes."""
    parts: List[str] = []
    buf: List[str] = []
    depth = 0
    in_quotes = False
    escape = False

    for ch in s:
        if escape:
            buf.append(ch)
            escape = False
            continue
        if ch == "\\":
            buf.append(ch)
            escape = True
            continue
        if ch == '"':
            in_quotes = not in_quotes
            buf.append(ch)
            continue

        if not in_quotes:
            if ch == "{":
                depth += 1
            elif ch == "}":
                depth = max(0, depth - 1)
            if ch == sep and depth == 0:
                part = "".join(buf).strip()
                if part:
                    parts.append(part)
                buf = []
                continue

        buf.append(ch)

    tail = "".join(buf).strip()
    if tail:
        parts.append(tail)
    return parts


def _strip_enclosing(value: str) -> str:
    value = value.strip()
    if value.endswith(","):
        value = value[:-1].rstrip()
    if len(value) >= 2 and ((value[0] == "{" and value[-1] == "}") or (value[0] == '"' and value[-1] == '"')):
        value = value[1:-1].strip()
    return value


@dataclass
class BibEntry:
    key: str
    entry_type: str
    fields: Dict[str, str]
    raw: str

    def get(self, field: str, default: str = "") -> str:
        return self.fields.get(field.lower(), default)


def _print_bib_error(msg: str, raw: str) -> None:
    sys.stderr.write("\n" + "=" * 88 + "\n")
    sys.stderr.write(msg.rstrip() + "\n")
    sys.stderr.write("-" * 88 + "\n")
    sys.stderr.write(raw.rstrip() + "\n")
    sys.stderr.write("=" * 88 + "\n")


def parse_bibtex_file(path: Path) -> List[BibEntry]:
    """Parse a BibTeX file into entries.

    Failure mode:
      If a specific record cannot be parsed, it is printed to stderr and skipped.
    """
    text = path.read_text(encoding="utf-8", errors="replace")
    n = len(text)
    i = 0

    entries: List[BibEntry] = []

    while True:
        at = text.find("@", i)
        if at == -1:
            break

        # Parse entry type
        j = at + 1
        while j < n and text[j].isspace():
            j += 1
        k = j
        while k < n and (text[k].isalnum() or text[k] in ["_", "-"]):
            k += 1
        entry_type = text[j:k].strip().lower()

        # Find opening brace/paren
        while k < n and text[k].isspace():
            k += 1
        if k >= n or text[k] not in "{(":
            i = at + 1
            continue
        open_ch = text[k]
        close_ch = "}" if open_ch == "{" else ")"
        k += 1

        # Find key (up to first comma)
        comma = text.find(",", k)
        if comma == -1:
            raw = text[at: min(n, at + 2000)]
            _print_bib_error("BibTeX parse error: could not find key comma.", raw)
            break
        key = text[k:comma].strip()

        # Find end of entry by matching braces
        depth = 1
        m = comma + 1
        in_quotes = False
        escape = False
        while m < n and depth > 0:
            ch = text[m]
            if escape:
                escape = False
            else:
                if ch == "\\":
                    escape = True
                elif ch == '"':
                    in_quotes = not in_quotes
                elif not in_quotes:
                    if ch == open_ch:
                        depth += 1
                    elif ch == close_ch:
                        depth -= 1
            m += 1

        if depth != 0:
            raw = text[at: min(n, at + 4000)]
            _print_bib_error("BibTeX parse error: unmatched braces/parentheses; skipping record.", raw)
            i = at + 1
            continue

        raw_entry = text[at:m]
        body = text[comma + 1 : m - 1].strip()

        try:
            fields: Dict[str, str] = {}
            for part in _split_top_level(body, sep=","):
                if "=" not in part:
                    continue
                name, val = part.split("=", 1)
                name = name.strip().lower()
                val = _strip_enclosing(val)
                fields[name] = val
            if not key:
                _print_bib_error("BibTeX parse warning: empty key; skipping record.", raw_entry)
            else:
                entries.append(BibEntry(key=key, entry_type=entry_type, fields=fields, raw=raw_entry))
        except Exception as ex:
            _print_bib_error(f"BibTeX parse error for record (key={key!r}): {ex}", raw_entry)

        i = m

    return entries


# ----------------------------
# Utilities: name matching
# ----------------------------


def _strip_latex(s: str) -> str:
    if s is None:
        return ""
    s = re.sub(r"[{}]", "", s)
    s = re.sub(r"\\[A-Za-z]+", "", s)
    s = s.replace("~", " ")
    return s


def _normalize_for_compare(s: str) -> str:
    s = _strip_latex(s)
    s = unicodedata.normalize("NFKD", s)
    s = "".join(ch for ch in s if not unicodedata.combining(ch))
    s = s.lower()
    s = re.sub(r"[\s\.,\-\,\'`\"]+", "", s)
    return s


def _normalize_tokens(s: str) -> List[str]:
    s = _strip_latex(s)
    s = unicodedata.normalize("NFKD", s)
    s = "".join(ch for ch in s if not unicodedata.combining(ch))
    s = s.replace("~", " ")
    s = re.sub(r"[^A-Za-z\s]", " ", s)
    return [t.lower() for t in s.split() if t.strip()]


def _parse_author_name(author: str) -> Tuple[str, str]:
    author = author.strip()
    if not author:
        return "", ""
    if "," in author:
        last, first = author.split(",", 1)
        return last.strip(), first.strip()
    parts = author.split()
    if len(parts) == 1:
        return parts[0], ""
    return parts[-1], " ".join(parts[:-1])


def author_matches_investigator(author: str, inv_first: str, inv_last: str) -> bool:
    last, first = _parse_author_name(author)
    if _normalize_for_compare(last) != _normalize_for_compare(inv_last):
        return False

    inv_tokens = _normalize_tokens(inv_first)
    inv_full = "".join(inv_tokens)
    inv_initial = inv_full[0] if inv_full else ""

    auth_tokens = _normalize_tokens(first)
    auth_full = "".join(auth_tokens)

    if not auth_tokens:
        return True

    is_initials = all(len(t) == 1 for t in auth_tokens)
    if is_initials:
        return bool(inv_initial) and inv_initial == auth_tokens[0][0]

    if inv_full and auth_full:
        if inv_full == auth_full:
            return True
        if len(inv_full) >= 3 and auth_full.startswith(inv_full):
            return True
        if len(auth_full) >= 3 and inv_full.startswith(auth_full):
            return True

    return False


def entry_investigators(entry: BibEntry, investigators: Sequence[Tuple[str, str]]) -> List[Tuple[str, str]]:
    author_field = entry.get("author", "")
    if not author_field:
        return []
    authors = re.split(r"\s+and\s+", author_field)
    matches: List[Tuple[str, str]] = []
    for inv_first, inv_last in investigators:
        for a in authors:
            if author_matches_investigator(a, inv_first, inv_last):
                matches.append((inv_first, inv_last))
                break
    return matches


def sort_investigators(investigators: Sequence[Tuple[str, str]]) -> List[Tuple[str, str]]:
    return sorted(
        investigators,
        key=lambda inv: (_normalize_for_compare(inv[1]), _normalize_for_compare(inv[0])),
    )


# ----------------------------
# Date inference
# ----------------------------


def guess_entry_date(entry: BibEntry, *, date_field: str = "auto") -> Optional[dt.date]:
    df = date_field.lower()

    def get_by_field(field: str) -> Optional[dt.date]:
        v = entry.get(field, "")
        if not v:
            return None
        return parse_iso_date(v)

    if df != "auto":
        d = get_by_field(df)
        return d

    # auto preference list
    for field in ["date", "year-month", "date-added", "year"]:
        if field == "year-month":
            y = entry.get("year", "").strip()
            mo = entry.get("month", "").strip()
            if y and mo:
                mo_norm = mo.lower().strip()
                month_map = {
                    "jan": 1,
                    "january": 1,
                    "feb": 2,
                    "february": 2,
                    "mar": 3,
                    "march": 3,
                    "apr": 4,
                    "april": 4,
                    "may": 5,
                    "jun": 6,
                    "june": 6,
                    "jul": 7,
                    "july": 7,
                    "aug": 8,
                    "august": 8,
                    "sep": 9,
                    "sept": 9,
                    "september": 9,
                    "oct": 10,
                    "october": 10,
                    "nov": 11,
                    "november": 11,
                    "dec": 12,
                    "december": 12,
                }
                mo_i: Optional[int]
                try:
                    mo_i = int(re.sub(r"\D", "", mo_norm)) if re.search(r"\d", mo_norm) else month_map.get(mo_norm)
                except ValueError:
                    mo_i = None
                if mo_i and y.isdigit():
                    return dt.date(int(y), mo_i, 15)
            continue

        d = get_by_field(field)
        if d is not None:
            return d

    y = entry.get("year", "").strip()
    if y.isdigit():
        return dt.date(int(y), 7, 1)

    return None


# ----------------------------
# Report data model
# ----------------------------


NUM_WORDS = {
    1: "one",
    2: "two",
    3: "three",
    4: "four",
    5: "five",
    6: "six",
    7: "seven",
    8: "eight",
    9: "nine",
    10: "ten",
    11: "eleven",
    12: "twelve",
    13: "thirteen",
    14: "fourteen",
    15: "fifteen",
    16: "sixteen",
    17: "seventeen",
    18: "eighteen",
    19: "nineteen",
    20: "twenty",
}


def suffix_for_year(year_index: int) -> str:
    word = NUM_WORDS.get(year_index)
    if word:
        return word
    return f"y{year_index:02d}"


def cite_command_for_suffix(suffix: str) -> str:
    return f"cite{suffix}"


def color_for_investigator_count(n: int) -> str:
    if n <= 1:
        return "black"
    if n == 2:
        return "blue"
    return "invOrange"


def compute_assignments(
    *,
    entries: Sequence[BibEntry],
    investigators: Sequence[Tuple[str, str]],
    periods: Sequence[Period],
    start_date: dt.date,
    end_date: dt.date,
    date_field: str,
) -> Tuple[
    Dict[str, dt.date],
    Dict[str, List[Tuple[str, str]]],
    Dict[str, int],
    Dict[Tuple[str, str], Dict[int, List[str]]],
]:
    """Compute mapping from investigators to period -> bib keys.

    Returns:
      entry_dates: key -> date
      entry_inv_matches: key -> list of investigators (first,last) on that paper
      entry_inv_count: key -> number of investigators on that paper
      inv_to_period_keys: inv -> period_index -> list of keys
    """

    entry_dates: Dict[str, dt.date] = {}
    entry_inv_matches: Dict[str, List[Tuple[str, str]]] = {}
    entry_inv_count: Dict[str, int] = {}

    for e in entries:
        try:
            d = guess_entry_date(e, date_field=date_field)
            if d is not None:
                entry_dates[e.key] = d
        except Exception as ex:
            _print_bib_error(f"Error inferring date for key={e.key!r}: {ex}", e.raw)

        try:
            matches = entry_investigators(e, investigators)
            entry_inv_matches[e.key] = matches
            entry_inv_count[e.key] = len(matches)
        except Exception as ex:
            _print_bib_error(f"Error matching investigators for key={e.key!r}: {ex}", e.raw)
            entry_inv_matches[e.key] = []
            entry_inv_count[e.key] = 0

    inv_to_period_keys: Dict[Tuple[str, str], Dict[int, List[str]]] = {
        inv: {p.index: [] for p in periods} for inv in investigators
    }

    for e in entries:
        try:
            d = entry_dates.get(e.key)
            if d is None:
                continue
            if d < start_date or d > end_date:
                continue

            period_idx: Optional[int] = None
            for p in periods:
                if p.start <= d < p.end:
                    period_idx = p.index
                    break
            if period_idx is None:
                continue

            matching_invs = entry_inv_matches.get(e.key, [])
            for inv in matching_invs:
                if inv in inv_to_period_keys:
                    inv_to_period_keys[inv][period_idx].append(e.key)
        except Exception as ex:
            _print_bib_error(f"Error processing record key={e.key!r}: {ex}", e.raw)
            continue

    # sort within each cell: date desc, then key
    for inv, per_map in inv_to_period_keys.items():
        for idx, keys in per_map.items():
            keys.sort(key=lambda k: (entry_dates.get(k, dt.date.min), k), reverse=True)

    return entry_dates, entry_inv_matches, entry_inv_count, inv_to_period_keys


# ----------------------------
# CSV output
# ----------------------------


def write_summary_csv(
    *,
    path: Path,
    investigators: Sequence[Tuple[str, str]],
    periods: Sequence[Period],
    inv_to_period_keys: Dict[Tuple[str, str], Dict[int, List[str]]],
) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    header = ["Investigator"] + [p.label for p in periods]
    with path.open("w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(header)
        for inv_first, inv_last in investigators:
            inv_name = f"{inv_first} {inv_last}".strip()
            row = [inv_name]
            per_map = inv_to_period_keys.get((inv_first, inv_last), {})
            for p in periods:
                keys = per_map.get(p.index, [])
                row.append("; ".join(keys))
            w.writerow(row)


# ----------------------------
# LaTeX generation
# ----------------------------


def latex_escape(s: str) -> str:
    """Escape a small set of LaTeX special chars for investigator names."""
    # Names typically do not include much, but do basic safety.
    repl = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
    }
    out = []
    for ch in s:
        out.append(repl.get(ch, ch))
    return "".join(out)


def latex_citations_for_keys(
    *,
    keys: Sequence[str],
    year_suffix: str,
    entry_inv_count: Dict[str, int],
) -> str:
    if not keys:
        return ""
    cite_cmd = cite_command_for_suffix(year_suffix)
    parts: List[str] = []
    for k in keys:
        c = color_for_investigator_count(entry_inv_count.get(k, 0))
        parts.append(rf"\textcolor{{{c}}}{{\{cite_cmd}{{{k}}}}}")
    return ", ".join(parts)


def generate_summary_longtable(
    *,
    investigators: Sequence[Tuple[str, str]],
    periods: Sequence[Period],
    inv_to_period_keys: Dict[Tuple[str, str], Dict[int, List[str]]],
    entry_inv_count: Dict[str, int],
) -> str:
    year_totals: List[int] = []
    for p in periods:
        keys: set[str] = set()
        for per_map in inv_to_period_keys.values():
            keys.update(per_map.get(p.index, []))
        year_totals.append(len(keys))

    investigator_col = r">{\raggedright\arraybackslash}p{1.2in}"
    year_col = r">{\raggedright\arraybackslash}p{1.0in}"
    colspec_str = "@{}" + investigator_col + "".join(year_col for _ in periods) + "@{}"

    lines: List[str] = []
    lines.append(r"\subsection*{Summary Table}")
    lines.append(r"\scriptsize")
    lines.append(r"\setlength{\LTpre}{0pt}")
    lines.append(r"\setlength{\LTpost}{0pt}")
    lines.append(r"\renewcommand{\arraystretch}{1.25}")
    lines.append(r"\setlength\tabcolsep{4pt}")

    lines.append(rf"\begin{{longtable}}{{{colspec_str}}}")
    lines.append(r"\toprule")
    header_cells = ["Investigator"] + [f"{p.label} ({n})" for p, n in zip(periods, year_totals)]
    header_row = " & ".join(rf"\textbf{{{latex_escape(c)}}}" for c in header_cells) + r"\\"
    lines.append(header_row)
    lines.append(r"\midrule")
    lines.append(r"\endfirsthead")

    lines.append(r"\toprule")
    lines.append(header_row)
    lines.append(r"\midrule")
    lines.append(r"\endhead")

    for inv_first, inv_last in investigators:
        inv_name = latex_escape(f"{inv_first} {inv_last}".strip())
        per_map = inv_to_period_keys.get((inv_first, inv_last), {})
        row_cells: List[str] = [inv_name]
        for p in periods:
            keys = per_map.get(p.index, [])
            year_suffix = suffix_for_year(p.index)
            cites = latex_citations_for_keys(keys=keys, year_suffix=year_suffix, entry_inv_count=entry_inv_count)
            row_cells.append(cites if cites else "")
        lines.append(" & ".join(row_cells) + r"\\")

    lines.append(r"\bottomrule")
    lines.append(r"\end{longtable}")
    lines.append(r"\normalsize")
    return "\n".join(lines)


def generate_tex(
    *,
    investigators: Sequence[Tuple[str, str]],
    periods: Sequence[Period],
    start_date: dt.date,
    end_date: dt.date,
    bib_resource_filename: str,
    inv_to_period_keys: Dict[Tuple[str, str], Dict[int, List[str]]],
    entry_dates: Dict[str, dt.date],
    entry_inv_count: Dict[str, int],
    include_bibliography: bool,
    bibliography_style: str,
) -> str:
    bib_base = Path(bib_resource_filename).stem

    # Precompute total pubs per investigator (unique keys across all periods)
    inv_total: Dict[Tuple[str, str], int] = {}
    for inv in investigators:
        per_map = inv_to_period_keys.get(inv, {})
        s: set[str] = set()
        for keys in per_map.values():
            s.update(keys)
        inv_total[inv] = len(s)

    # Preamble
    lines: List[str] = []
    lines.append(r"\documentclass[11pt]{article}")
    lines.append(r"\usepackage[margin=1in]{geometry}")
    lines.append(r"\usepackage[hidelinks]{hyperref}")
    lines.append(r"\usepackage{xcolor}")
    lines.append(r"\definecolor{invOrange}{RGB}{255,130,0}")
    lines.append(r"\usepackage{array}")
    lines.append(r"\usepackage{longtable}")
    lines.append(r"""
% -------------------------------------------------------------------------------   
% Tables
% -------------------------------------------------------------------------------   
\usepackage{booktabs}

% this improves table legibility
\AtBeginDocument{%
\heavyrulewidth=.08em
\lightrulewidth=.05em
\cmidrulewidth=.03em
\belowrulesep=.65ex
\belowbottomsep=0pt
\aboverulesep=.4ex
\abovetopsep=0pt
\cmidrulesep=\doublerulesep
\cmidrulekern=.5em
\defaultaddspace=.5em
}
% -------------------------------------------------------------------------------   
""")
    lines.append(r"\usepackage[numbers,sort&compress]{natbib}")
    lines.append(r"\usepackage{multibib}")

    # multibib declarations: one per year
    for p in periods:
        suffix = suffix_for_year(p.index)
        lines.append(rf"\newcites{{{suffix}}}{{{p.label}}}")

    lines.append("")
    lines.append(r"\begin{document}")
    lines.append("")
    lines.append(r"\section*{IRG1 Investigator Publications by Project Year}")
    lines.append(rf"\noindent Project date range: {start_date.isoformat()} $\to$ {end_date.isoformat()}\\")
    lines.append(r"\noindent Citation colors: black = 1 investigator on the paper, \textcolor{blue}{blue = 2}, \textcolor{invOrange}{orange = 3+}.")
    lines.append("")

    # Summary table at top
    lines.append(generate_summary_longtable(
        investigators=investigators,
        periods=periods,
        inv_to_period_keys=inv_to_period_keys,
        entry_inv_count=entry_inv_count,
    ))
    lines.append("")

    # Investigator sections
    for inv_first, inv_last in investigators:
        inv_name = latex_escape(f"{inv_first} {inv_last}".strip())
        total = inv_total.get((inv_first, inv_last), 0)
        lines.append(rf"\section*{{{inv_name} ({total})}}")
        lines.append(r"\begin{itemize}")

        per_map = inv_to_period_keys.get((inv_first, inv_last), {})
        for p in periods:
            keys = per_map.get(p.index, [])
            suffix = suffix_for_year(p.index)
            cites = latex_citations_for_keys(keys=keys, year_suffix=suffix, entry_inv_count=entry_inv_count)
            if not cites:
                cites = "None"
            lines.append(rf"\item \textbf{{{p.label}:}} {cites}")

        lines.append(r"\end{itemize}")
        lines.append("")

    # Yearly bibliographies
    if include_bibliography:
        lines.append(r"\newpage")
        lines.append(r"\section*{Yearly Bibliographies}")
        lines.append("")
        for p in periods:
            suffix = suffix_for_year(p.index)
            lines.append(rf"\bibliographystyle{suffix}{{{bibliography_style}}}")
            lines.append(rf"\bibliography{suffix}{{{bib_base}}}")
            lines.append("")

    lines.append(r"\end{document}")
    return "\n".join(lines) + "\n"


def _strip_tex_braces(s: str) -> str:
    prev = None
    while prev != s:
        prev = s
        s = s.replace("{", "").replace("}", "")
    return s


def _latex_to_plain_text(s: str) -> str:
    """Best-effort cleanup of BibTeX-generated .bbl text for Word output."""
    s = s.replace("\n", " ")
    s = s.replace("~", " ")

    s = re.sub(r"\\newblock\s*", " ", s)
    s = re.sub(r"\\url\{([^}]*)\}", r"\1", s)
    s = re.sub(r"\\doi\{([^}]*)\}", r"DOI: \1", s)
    s = re.sub(r"\\href\{[^}]*\}\{([^}]*)\}", r"\1", s)
    s = re.sub(r"\\emph\{([^}]*)\}", r"\1", s)
    s = re.sub(r"\\textit\{([^}]*)\}", r"\1", s)
    s = re.sub(r"\\textbf\{([^}]*)\}", r"\1", s)
    s = re.sub(r"\\textsc\{([^}]*)\}", r"\1", s)
    s = re.sub(r"\\enquote\{([^}]*)\}", r"\1", s)
    s = re.sub(r"\\[a-zA-Z]+\*?\s*", " ", s)

    s = _strip_tex_braces(s)
    s = s.replace("``", '"').replace("''", '"')
    s = s.replace(r"\&", "&").replace(r"\%", "%").replace(r"\_", "_")
    s = s.replace(r"\$", "$").replace(r"\#", "#")
    s = html.unescape(s)
    s = re.sub(r"\s+([,.;:])", r"\1", s)
    s = re.sub(r"\(\s+", "(", s)
    s = re.sub(r"\s+\)", ")", s)
    s = re.sub(r"\s{2,}", " ", s)
    return s.strip()


def parse_bbl_file(path: Path) -> List[Tuple[str, str]]:
    """Return [(bibkey, rendered_text), ...] from a BibTeX-generated .bbl file."""
    text = path.read_text(encoding="utf-8", errors="replace")
    entries: List[Tuple[str, str]] = []

    pattern = re.compile(r"\\bibitem(?:\[[^\]]*\])?\{([^}]+)\}")
    matches = list(pattern.finditer(text))
    for i, m in enumerate(matches):
        key = m.group(1).strip()
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        body = text[start:end]
        body = re.sub(r"\\end\{thebibliography\}.*$", "", body, flags=re.DOTALL)
        rendered = _latex_to_plain_text(body)
        if rendered:
            entries.append((key, rendered))
    return entries


def write_docx_from_bbl(*, path: Path, periods: Sequence[Period], tex_dir: Path) -> None:
    if Document is None:
        raise RuntimeError("python-docx is not available. Install it with: pip install python-docx")

    doc = Document()
    doc.add_heading("Yearly Bibliographies", level=1)

    for p in periods:
        suffix = suffix_for_year(p.index)
        bbl_path = tex_dir / f"{suffix}.bbl"
        doc.add_heading(p.label, level=2)

        if not bbl_path.exists():
            doc.add_paragraph(f"Missing bibliography file: {bbl_path.name}")
            continue

        items = parse_bbl_file(bbl_path)
        if not items:
            doc.add_paragraph("None")
            continue

        for _, rendered in items:
            doc.add_paragraph(rendered, style="List Number")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


# ----------------------------
# LaTeX compilation
# ----------------------------


def compile_latex(tex_path: Path, *, year_suffixes: Sequence[str], include_bibliography: bool) -> None:
    """Compile LaTeX using pdflatex + bibtex (multibib) + pdflatex twice."""
    workdir = tex_path.parent

    bibtex_exe = shutil.which("bibtex")
    if bibtex_exe is None:
        # In some environments (including some minimal TeXLive installs), the bibtex symlink can be missing.
        bibtex_exe = shutil.which("bibtex.original") or "bibtex.original"

    # Clean old auxiliary and bibliography artifacts for deterministic builds
    stem = tex_path.stem
    for ext in ["aux", "bbl", "blg", "log", "out", "toc"]:
        try:
            (workdir / f"{stem}.{ext}").unlink()
        except FileNotFoundError:
            pass
    for suf in year_suffixes:
        for ext in ["aux", "bbl", "blg"]:
            try:
                (workdir / f"{suf}.{ext}").unlink()
            except FileNotFoundError:
                pass

    def run(cmd: Sequence[str]) -> None:
        proc = subprocess.run(
            cmd,
            cwd=str(workdir),
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
        )
        if proc.returncode != 0:
            sys.stderr.write(proc.stdout)
            raise RuntimeError(f"Command failed: {' '.join(cmd)}")

    run(["pdflatex", "-interaction=nonstopmode", "-halt-on-error", tex_path.name])

    if include_bibliography:
        # multibib creates aux files named exactly '<suffix>.aux'
        for suf in year_suffixes:
            aux = workdir / f"{suf}.aux"
            if not aux.exists():
                continue
            # Only run bibtex if bibdata appears; otherwise bibtex fails noisily.
            try:
                content = aux.read_text(encoding="utf-8", errors="replace")
            except Exception:
                content = ""
            if "\\bibdata" not in content:
                continue
            if "\\citation" not in content:
                # No citations for this year; skip bibtex to avoid nonzero exit status.
                continue
            run([bibtex_exe, suf])

    run(["pdflatex", "-interaction=nonstopmode", "-halt-on-error", tex_path.name])
    run(["pdflatex", "-interaction=nonstopmode", "-halt-on-error", tex_path.name])


# ----------------------------
# CLI
# ----------------------------


def load_investigators(csv_path: Path) -> List[Tuple[str, str]]:
    """Load investigators from CSV.

    Required columns (case-insensitive):
      - Firstname (or: first, given, givenname, first_name, given_name)
      - Lastname  (or: last, family, familyname, surname, last_name, family_name)

    Returns list of (first,last) in file order.
    """
    with csv_path.open(newline="", encoding="utf-8", errors="replace") as f:
        reader = csv.DictReader(f)
        if reader.fieldnames is None:
            raise ValueError("Investigators CSV has no header row.")

        headers = [h.strip() for h in reader.fieldnames]
        header_map = {h.lower().strip(): h for h in headers}

        def find_col(candidates: Sequence[str]) -> Optional[str]:
            for c in candidates:
                if c in header_map:
                    return header_map[c]
            return None

        first_col = find_col(["firstname", "first", "given", "givenname", "first_name", "given_name"])
        last_col = find_col(["lastname", "last", "family", "familyname", "surname", "last_name", "family_name"])
        if not first_col or not last_col:
            raise ValueError(f"Investigators CSV must contain first/last name columns. Found: {headers}")

        out: List[Tuple[str, str]] = []
        for row in reader:
            first = (row.get(first_col) or "").strip()
            last = (row.get(last_col) or "").strip()
            if not first and not last:
                continue
            out.append((first, last))
        return out


def main(argv: Optional[Sequence[str]] = None) -> int:
    p = argparse.ArgumentParser(
        description="Generate a LaTeX report of citations per investigator per project-year (multibib)."
    )
    p.add_argument("bib", type=Path, help="Input .bib file")
    p.add_argument("investigators", type=Path, help="Investigators CSV")
    p.add_argument("--start-date", required=True, help="YYYY-MM-DD")
    p.add_argument("--end-date", required=True, help="YYYY-MM-DD")
    p.add_argument("--out", dest="out_tex", type=Path, default=Path("report.tex"), help="Output .tex file")
    p.add_argument(
        "--summary-csv",
        dest="summary_csv",
        type=Path,
        default=None,
        help="Optional path for summary CSV (default: summary_pubs.csv next to --out).",
    )
    p.add_argument(
        "--date-field",
        default="auto",
        help=(
            "Which BibTeX field to use for dates: auto (default), date, date-added, year, month, etc. "
            "If not auto, no fallback is used."
        ),
    )
    p.add_argument(
        "--bibliography-style",
        default="plain",
        help="BibTeX bibliography style to use for each year (default: plain).",
    )
    p.add_argument(
        "--no-bibliography",
        action="store_true",
        help=r"Do not include per-year bibliographies in the output TeX.",
    )
    p.add_argument(
        "--compile",
        action="store_true",
        help="Compile to PDF with pdflatex + bibtex (multibib).",
    )
    p.add_argument(
        "--docx",
        action="store_true",
        help="Also write a .docx file containing only the rendered yearly bibliographies from the generated .bbl files.",
    )
    p.add_argument(
        "--docx-out",
        dest="docx_out",
        type=Path,
        default=None,
        help="Optional path for .docx output (default: same stem as --out, with .docx suffix).",
    )

    args = p.parse_args(argv)

    start = parse_iso_date(args.start_date)
    end = parse_iso_date(args.end_date)
    if start is None or end is None:
        raise ValueError("--start-date and --end-date must be valid ISO dates (YYYY-MM-DD).")
    if end < start:
        raise ValueError("--end-date must be >= --start-date")

    investigators = load_investigators(args.investigators)
    if not investigators:
        raise ValueError("No investigators found in CSV")
    investigators = sort_investigators(investigators)

    entries = parse_bibtex_file(args.bib)
    periods = build_periods(start, end)

    entry_dates, entry_inv_matches, entry_inv_count, inv_to_period_keys = compute_assignments(
        entries=entries,
        investigators=investigators,
        periods=periods,
        start_date=start,
        end_date=end,
        date_field=args.date_field,
    )

    out_tex = args.out_tex
    out_tex.parent.mkdir(parents=True, exist_ok=True)

    # Copy bib file next to TeX output for compilation portability.
    bib_dest = out_tex.parent / args.bib.name
    if args.bib.resolve() != bib_dest.resolve():
        shutil.copy2(args.bib, bib_dest)

    # Write summary CSV
    summary_csv_path = args.summary_csv or (out_tex.parent / "summary_pubs.csv")
    write_summary_csv(
        path=summary_csv_path,
        investigators=investigators,
        periods=periods,
        inv_to_period_keys=inv_to_period_keys,
    )

    # Write TeX
    tex = generate_tex(
        investigators=investigators,
        periods=periods,
        start_date=start,
        end_date=end,
        bib_resource_filename=bib_dest.name,
        inv_to_period_keys=inv_to_period_keys,
        entry_dates=entry_dates,
        entry_inv_count=entry_inv_count,
        include_bibliography=not args.no_bibliography,
        bibliography_style=args.bibliography_style,
    )
    out_tex.write_text(tex, encoding="utf-8")

    if args.compile:
        year_suffixes = [suffix_for_year(p.index) for p in periods]
        compile_latex(out_tex, year_suffixes=year_suffixes, include_bibliography=not args.no_bibliography)

    if args.docx:
        if args.no_bibliography:
            raise RuntimeError("--docx requires yearly bibliographies, so it cannot be used with --no-bibliography.")
        if not args.compile:
            raise RuntimeError("--docx from rendered .bbl files requires --compile so the .bbl files exist.")
        docx_path = args.docx_out or out_tex.with_suffix(".docx")
        write_docx_from_bbl(path=docx_path, periods=periods, tex_dir=out_tex.parent)

    return 0

